"""视频文案提取工具 —— 支持任意平台视频链接。"""
import io
import os
import tempfile
import threading
import uuid
from flask import Flask, render_template, request, jsonify, send_file

from analyzer import correct_transcript
from scraper import fetch_xhs_comments, fetch_note
from video import download_video, extract_audio, transcribe, get_video_title, get_comments

app = Flask(__name__)

UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "xhs_uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

tasks = {}

# IP 使用次数记录 {ip: count}
ip_usage = {}
BATCH_LIMIT = 10


def _get_ip():
    """获取真实客户端 IP（兼容 Railway 代理）。"""
    return (request.headers.get("X-Forwarded-For") or request.remote_addr or "").split(",")[0].strip()


def _check_batch_limit():
    """检查 IP 是否超出批量提取限额，返回 (ok, used, limit)。"""
    ip = _get_ip()
    used = ip_usage.get(ip, 0)
    return used < BATCH_LIMIT, used, BATCH_LIMIT


def _inc_batch_usage():
    ip = _get_ip()
    ip_usage[ip] = ip_usage.get(ip, 0) + 1


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/transcribe_link", methods=["POST"])
def transcribe_link():
    data = request.get_json(force=True)
    link = (data.get("link") or "").strip()
    cookie = (data.get("cookie") or "").strip() or None

    if not link:
        return jsonify(ok=False, error="请输入链接")

    task_id = str(uuid.uuid4())
    tasks[task_id] = {"status": "processing", "result": None, "error": None}

    def run():
        try:
            title = _get_title(link, cookie)
            video_path = download_video(link, cookie=cookie)
            audio_path = extract_audio(video_path)
            result = transcribe(audio_path)
            if title:
                result["title"] = title
            _cleanup(video_path, audio_path)
            tasks[task_id]["result"] = result
            tasks[task_id]["status"] = "done"
        except Exception as e:
            tasks[task_id]["error"] = str(e)
            tasks[task_id]["status"] = "error"

    threading.Thread(target=run, daemon=True).start()
    return jsonify(ok=True, task_id=task_id)


@app.route("/transcribe_upload", methods=["POST"])
def transcribe_upload():
    file = request.files.get("video")
    if not file:
        return jsonify(ok=False, error="请上传视频文件")

    video_path = os.path.join(UPLOAD_DIR, file.filename)
    file.save(video_path)

    task_id = str(uuid.uuid4())
    tasks[task_id] = {"status": "processing", "result": None, "error": None}

    def run():
        try:
            audio_path = extract_audio(video_path)
            result = transcribe(audio_path)
            _cleanup(video_path, audio_path)
            tasks[task_id]["result"] = result
            tasks[task_id]["status"] = "done"
        except Exception as e:
            tasks[task_id]["error"] = str(e)
            tasks[task_id]["status"] = "error"

    threading.Thread(target=run, daemon=True).start()
    return jsonify(ok=True, task_id=task_id)


@app.route("/task/<task_id>")
def get_task(task_id):
    t = tasks.get(task_id)
    if not t:
        return jsonify(ok=False, error="任务不存在")
    if t["status"] == "processing":
        return jsonify(ok=True, status="processing")
    if t["status"] == "error":
        tasks.pop(task_id, None)
        return jsonify(ok=False, error=t["error"])
    result = t["result"]
    tasks.pop(task_id, None)
    return jsonify(ok=True, status="done", **result)


@app.route("/correct", methods=["POST"])
def do_correct():
    data = request.get_json(force=True)
    content = (data.get("content") or "").strip()
    api_key = (data.get("api_key") or "").strip() or None
    if not content:
        return jsonify(ok=False, error="逐字稿为空，无法纠错")
    try:
        corrected = correct_transcript(content, api_key=api_key)
        return jsonify(ok=True, content=corrected)
    except Exception as e:
        return jsonify(ok=False, error=str(e))



@app.route("/quota")
def quota():
    ok, used, limit = _check_batch_limit()
    return jsonify(ok=True, used=used, limit=limit, remaining=limit-used)


@app.route("/batch_transcribe", methods=["POST"])
def batch_transcribe():
    data = request.get_json(force=True)
    import re
    raw_lines = (data.get("links") or "").strip().splitlines()
    cookie_xhs = (data.get("cookie_xhs") or "").strip() or None
    cookie_douyin = (data.get("cookie_douyin") or "").strip() or None
    # 只提取包含 URL 的行，忽略纯文字描述行
    links = []
    for line in raw_lines:
        m = re.search(r"https?://[^\s，,]+", line)
        if m:
            links.append(m.group(0))

    if not links:
        return jsonify(ok=False, error="请输入至少一个链接")
    if len(links) > 10:
        return jsonify(ok=False, error="最多同时处理 10 个链接")

    ok, used, limit = _check_batch_limit()
    if not ok:
        return jsonify(ok=False, error="QUOTA_EXCEEDED", used=used, limit=limit)
    _inc_batch_usage()

    task_ids = []
    for link in links:
        task_id = str(uuid.uuid4())
        tasks[task_id] = {"status": "processing", "result": None, "error": None, "link": link}

        def _cookie_for(lnk):
            if "xiaohongshu" in lnk or "xhslink" in lnk:
                return cookie_xhs
            if "douyin" in lnk or "iesdouyin" in lnk:
                return cookie_douyin
            return None

        def run(tid=task_id, lnk=link):
            try:
                cookie = _cookie_for(lnk)
                title = _get_title(lnk, cookie=cookie)
                video_path = download_video(lnk, cookie=cookie)
                audio_path = extract_audio(video_path)
                result = transcribe(audio_path)
                if title:
                    result["title"] = title
                result["link"] = lnk

                _cleanup(video_path, audio_path)
                tasks[tid]["result"] = result
                tasks[tid]["status"] = "done"
            except Exception as e:
                tasks[tid]["error"] = str(e)
                tasks[tid]["status"] = "error"

        threading.Thread(target=run, daemon=True).start()
        task_ids.append(task_id)

    return jsonify(ok=True, task_ids=task_ids)


@app.route("/comments", methods=["POST"])
def do_comments():
    data = request.get_json(force=True)
    link = (data.get("link") or "").strip()
    cookie = (data.get("cookie") or "").strip() or None
    if not link:
        return jsonify(ok=False, error="请输入链接")
    try:
        comments = None
        # 小红书用专用抓取
        if "xiaohongshu" in link or "xhslink" in link:
            comments = fetch_xhs_comments(link, cookie=cookie)
        # 其他平台用 yt-dlp
        if not comments:
            comments = get_comments(link, cookie=cookie)
        if not comments:
            return jsonify(ok=False, error="未能提取到评论，可能需要在设置中填写 Cookie")
        return jsonify(ok=True, comments=comments)
    except Exception as e:
        return jsonify(ok=False, error=str(e))


@app.route("/export_docx", methods=["POST"])
def export_docx():
    data = request.get_json(force=True)
    title = (data.get("title") or "").strip()
    content = (data.get("content") or "").strip()
    if not content:
        return jsonify(ok=False, error="逐字稿为空，无法导出")

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for para_text in content.split("\n"):
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = (title[:30] if title else "逐字稿") + ".docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/export_batch_docx", methods=["POST"])
def export_batch_docx():
    data = request.get_json(force=True)
    items = data.get("items") or []
    if not items:
        return jsonify(ok=False, error="没有可导出的内容")

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()
    for i, item in enumerate(items):
        title = (item.get("title") or "").strip() or f"第 {i+1} 篇"
        content = (item.get("text") or "").strip()

        # 标题
        h = doc.add_heading(title, level=1)
        # 正文
        for line in content.split("\n"):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(12)

        # 分隔线（非最后一篇）
        if i < len(items) - 1:
            sep = doc.add_paragraph("─" * 30)
            sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in sep.runs:
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"批量提取_{len(items)}篇.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def _get_title(link, cookie=None):
    """多种方式尝试提取标题：yt-dlp → 小红书抓取器。"""
    # 1. yt-dlp
    title = get_video_title(link, cookie=cookie)
    if title:
        return title
    # 2. 小红书抓取器
    if "xiaohongshu" in link or "xhslink" in link:
        try:
            note = fetch_note(link, cookie=cookie)
            if note.get("title"):
                return note["title"]
        except Exception:
            pass
    return None


def _cleanup(*paths):
    for p in paths:
        try:
            os.remove(p)
        except OSError:
            pass


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=os.environ.get("FLASK_DEBUG") == "1")
